from http.client import OK
from os import access
from flask import Flask, render_template, request, redirect, url_for, session,abort

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
import docx
from docx import Document
from docx.shared import Inches
import secrets
import requests
import pandas as pd
from webdriver_manager.chrome import ChromeDriverManager
from flask_cors import CORS
from pymongo import MongoClient
from flask_mail import Mail,Message


chrome_options = Options()
chrome_options.add_argument("--headless")

def get_page(keyword):
    driver=webdriver.Chrome(ChromeDriverManager().install(),options=chrome_options)
    URL=driver.get('https://pubchem.ncbi.nlm.nih.gov/')
    search_key_word=keyword
    search=driver.find_element_by_xpath("//div[@class='flex-grow-1 p-xsm-left p-xsm-top p-xsm-bottom b-right b-light']/input").send_keys(search_key_word)
    find=driver.find_element_by_xpath("//button[@data-action='search-button']")
    find.click()
    result_loaded = EC.presence_of_element_located((By.XPATH,'//div[@id="featured-results"]//div[@class="f-medium p-sm-top p-sm-bottom f-1125"]/a'))
    WebDriverWait(driver, 3).until(result_loaded)
    best_match= driver.find_element_by_xpath("//div[@id='featured-results']//div[@class='f-medium p-sm-top p-sm-bottom f-1125']/a").click()
    
    try:
        ATC_codes_loaded = EC.presence_of_element_located((By.XPATH,"//*[@id='ATC-Code']"))
        WebDriverWait(driver,3).until(ATC_codes_loaded)
        ATC_codes=driver.find_element_by_xpath("//*[@id='ATC-Code']")
        list_atc_codes=[]
        for i in ATC_codes.find_elements_by_xpath("//a[@data-label='Content Link: ATC-Code']"):
            if(i.get_property("target")!="_self" and i.get_property("rel")!="follow"):
                list_atc_codes.append(i.find_element_by_xpath('..').text)
                list_atc_codes.append("")
        list_atc_codes
    except:
        list_atc_codes=" "
    
    try:
        IUPAC_name=driver.find_element_by_xpath("//section[@id='IUPAC-Name']//p")
        iupac=IUPAC_name.text
    except NoSuchElementException:
        iupac=" "
    
    try:
        physical_description=driver.find_element_by_xpath("//section[@id='Physical-Description']//p")
        phy_desc=physical_description.text
    except NoSuchElementException:
        phy_desc=" "
    
    try:
        phy_desc_reviews_find=driver.find_element_by_xpath("//section[@id='Physical-Description']//div[@class='peer-review f-0875 f-lh-115 italic f-gray regular-links breakword'][1]/*")
        phy_desc_review=phy_desc_reviews_find.text
    except NoSuchElementException:
        phy_desc_review=" "

    try:
        solubility_find_1=driver.find_element_by_xpath("//section[@id='Solubility']//*[contains(text(), 'CAMEO Chemicals')]/../../../preceding::p[1]")
        solubility_1=solubility_find_1.text
        solubility_1
    except NoSuchElementException:
        solubility_find_1=driver.find_element_by_xpath("//section[@id='Solubility']//div[@class='section-content-item'][1]/p")
        solubility_1=solubility_find_1.text
    except NoSuchElementException:
        solubility_1=" "

    try:
        solubility_find_1_review=driver.find_element_by_xpath("//section[@id='Solubility']//*[contains(text(), 'CAMEO Chemicals')]/../../../preceding::div[1]")
        solubility_1_review=solubility_find_1_review.text
        solubility_1_review
    except NoSuchElementException:
        solubility_1_review=" "

    try:
        melting_point_find=driver.find_element_by_xpath("//*[@id='Melting-Point']//div[2]//div/p")
        melting_point=melting_point_find.text
    except NoSuchElementException:
        melting_point=" "

    try:
        optical_rotation_find= driver.find_element_by_xpath("//section[@id='Optical-Rotation']/div[2]/div[1]/p")
        optical_rotation=optical_rotation_find.text
    except NoSuchElementException:
        optical_rotation=" "

    try:
        optical_rotation_find_review= driver.find_element_by_xpath("//section[@id='Optical-Rotation']/div[2]/div[1]/div")
        optical_rotation_review= optical_rotation_find_review.text
    except NoSuchElementException:
        optical_rotation_review=" "

    try:
        decomposition_find= driver.find_element_by_xpath("//section[@id='Decomposition']/div[2]/div/p")
        decomposition=decomposition_find.text
    except NoSuchElementException:
        decomposition=" "
    
    try:
        decomposition_review_find=driver.find_element_by_xpath("//section[@id='Decomposition']/div[2]/div/div")
        decomposition_review=decomposition_review_find.text
    except NoSuchElementException :
        decomposition_review=" "

    try:
        shelf_life_find=driver.find_element_by_xpath("//section[@id='Stability-Shelf-Life']/div[2]/div/p")
        shelf_life=shelf_life_find.text
    except NoSuchElementException:
        shelf_life=" "

    try:
        shelf_life_find_review=driver.find_element_by_xpath("//section[@id='Stability-Shelf-Life']/div[2]/div//div[@class='peer-review f-0875 f-lh-115 italic f-gray regular-links breakword']/*")
        shelf_life_review=shelf_life_find_review.text
    except NoSuchElementException:
        shelf_life_review=" "

    try:
        stability_find=driver.find_element_by_xpath("//section[@id='Stability-Shelf-Life']/div[2]/div[3]/p")
        stability=stability_find.text
    except NoSuchElementException:
        stability=" "

    try:
        molecular_formula_find=driver.find_element_by_xpath("//section[@id='Molecular-Formula']/div[2]/div/p")
        molecular_formula=molecular_formula_find.text
    except NoSuchElementException:
        molecular_formula=" "

    try:
        CAS_number_find=driver.find_element_by_xpath("//section[@id='CAS']/div[2]/div[1]/p")
        CAS_number=CAS_number_find.text
    except NoSuchElementException:
        CAS_number=" "

    try:
        molecular_weight_find=driver.find_element_by_xpath("//section[@id='Computed-Properties']/div[2]/div/div[1]/table/tbody/tr[1]/td[2]")
        molecular_weight=molecular_weight_find.text
    except NoSuchElementException:
        molecular_weight=" "
    try:
        chemical_tructure=driver.find_element_by_xpath("//*[@id='2D-Structure']/div[2]/div[1]/div/div[2]/table/tr/td/div/div[1]/div/div/img").get_attribute("src")
        chemical_tructure
    except NoSuchElementException:
        molecular_weight=" "
    
    with open("static/images/"+str(search_key_word)+'.jpg', 'wb') as handle:

        response = requests.get(chemical_tructure, stream=True)

        if not response.ok:
            print(response)

        for block in response.iter_content(1024):
            if not block:
                break

            handle.write(block)

    doc = docx.Document()
    head=doc.add_heading('2.5.1 Product Devlopment Rationale')
    drug_name=keyword
    doc_para = doc.add_paragraph('Under the Anatomical Therapeutic Chemical (ATC) classification system'+drug_name+' are classified under the following headings  (« WHOCC - ATC/DDD Index » s. d.): ' )
    doc.add_paragraph(list_atc_codes)
    doc_para = doc.add_paragraph('In the following sub-modules, an integrated and critical assessment on pharmacodynamics, pharmacokinetics, efficacy and safety is presented in order to support Marketing Authorisation application for '+drug_name+'under the legal status of a Generic medicinal product, according to Article 10 (1) of the Directive 2001/83/EC, as last amended.' )
    doc_para= doc.add_paragraph('The clinical documentation has been obtained using standard text and reference books as well as computer-based searches of the scientific literature.')
    doc.add_heading('2.5.2 Overview of Biopharmaceutics')
    doc.add_paragraph('The chemical name of '+ search_key_word + ' is '+iupac+' (PubChem s.d.).')
    doc_para=doc.add_paragraph(drug_name+'has the following experimental properties (PubChem s. d.) :')
    doc.add_paragraph('physical description : '+phy_desc+'.'+"\n"+phy_desc_review)
    #write review
    doc.add_paragraph('solubility : '+solubility_1+"\n"+solubility_1_review)
    doc.add_paragraph('Melting point : '+melting_point)
    doc.add_paragraph('Optical rotation : '+optical_rotation+"\n"+optical_rotation_review)
    doc.add_paragraph('Decomposition : '+decomposition+"\n"+decomposition_review)
    doc_para=doc.add_paragraph('Stability/Shelf life : ')
    doc.add_paragraph(shelf_life+"\n"+shelf_life_review)
    doc.add_paragraph(stability)
    doc_para=doc.add_paragraph('The table below (Table 1) summarises chemical information on (PubChem s. d.).')
    doc_para=doc.add_paragraph('Table 1: chemical information')
    doc_para.alignment =1
    doc_para=doc.add_paragraph('____________________________________________________________________________________________________')
    doc.add_picture("static/images/"+rf"{search_key_word}.jpg")
    doc_para=doc.add_paragraph('____________________________________________________________________________________________________')
    doc.add_paragraph('Molecular formula : '+molecular_formula)
    doc.add_paragraph('Molecular weight : '+molecular_weight)
    doc.add_paragraph('CAS Registry Number : '+CAS_number)
    doc.save("static/articles/"+search_key_word +".docx")

    return {"ATC-Code":list_atc_codes,
            "iupac":iupac,"phy_desc":phy_desc,
            "phy_desc_review":phy_desc_review,
            "solubility_1":solubility_1,
            "solubility_1_review":solubility_1_review,
            "melting_point":melting_point,
            "optical_rotation":optical_rotation,
            "optical_rotation_review":optical_rotation_review,
            "decomposition":decomposition,
            "decomposition_review":decomposition_review,
            "shelf_life":shelf_life,
            "shelf_life_review":shelf_life_review,
            "stability":stability,
            "molecular_formula":molecular_formula,
            "CAS_number":CAS_number,
            "molecular_weight":molecular_weight,
            }

def get_database():
    
    # Provide the mongodb atlas url to connect python to mongodb using pymongo
    CONNECTION_STRING = "mongodb://localhost:27017"

    # Create a connection using MongoClient. You can import MongoClient or use pymongo.MongoClient
    client = MongoClient(CONNECTION_STRING)

    # Create the database for our example (we will use the same database throughout the tutorial
    return client


# Intialize Flask
app = Flask(__name__)
# Change this to your secret key (can be anything, it's for extra protection)
app.secret_key = '%3&}QWkq+>y7<pqZ'

CORS(app)


app.config['MAIL_SERVER']='smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USERNAME'] = 'medissearch@gmail.com'
app.config['MAIL_PASSWORD'] = 'yourpassword'
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False

mail = Mail(app)


def getUsers():    
    db = get_database()
    return db["medis"]["users"]



@app.route('/get_keys',methods=['GET'])
def get_features():
    molucule = request.args.get('molucule')
    return get_page(molucule)

@app.route('/login',methods=['POST'])
def login_user():
    data = request.get_json(force=True)
    email =data["email"]
    password = data["password"]
    user = list(getUsers().find({ "email": email }))
    if(len(user)>0):
        if(password == user[0]["password"]):
            return ({"email":user[0]["email"],"role":user[0]["role"],"username":user[0]["username"]},200)
        else:
            return ("Please verify your password",404)
    else:
        return ("No account found with this email",404)

@app.route('/users',methods=['GET','POST','DELETE'])
def Users():
    if(request.method=="GET"):
        try:
            users = list(getUsers().find({},{"_id": 0, "email": 1, "username": 1,"role":1}))
            return ({"users":users},200)
        except:
            return ("Error getting users",500)
    elif(request.method=="POST"):
        try:
            data = request.get_json(force=True)
            password = secrets.token_hex(nbytes=10)
            getUsers().insert_one({"username":data["username"],"email":data["email"],"password":password,"role":"user"})

            msg = Message('Your credentials at MediSearch tool', sender =   'admin@medis.tn', recipients = [data["email"]])
            msg.body = "Hello you are now registred at MediSearch tool and you can login using your e-mail and the following password : "+password 
            mail.send(msg)

            users = list(getUsers().find({},{"_id": 0, "email": 1, "username": 1,"role":1}))
            return ({"users":users},200)
        except Exception as e :
            print(e)
            return ("Error Creating user",500)
    elif(request.method=="DELETE"):
        try:
            email = request.args.get('email')
            getUsers().delete_one({"email":email})
            users = list(getUsers().find({},{"_id": 0, "email": 1, "username": 1,"role":1}))
            return ({"users":users},200)
        except:
            return ("Error Deleting user",500)


if __name__ == "__main__":

    app.run(host="0.0.0.0")
